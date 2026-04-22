<<<<<<< HEAD
import json
import os
import uuid
import sqlite3
import urllib.parse
import time
import re
from datetime import datetime

import requests
from flask import Flask, jsonify, render_template, request, send_from_directory
from openai import OpenAI

app = Flask(__name__)

BASE_DIR = os.path.dirname(__file__)
DB_PATH = os.path.join(BASE_DIR, "books.db")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OFFICIAL_DIR = os.path.join(UPLOAD_DIR, "official")
ALLOWED_UPLOAD_EXTENSIONS = {".txt", ".docx", ".pdf"}
WORD_PATTERN = re.compile(r"^[A-Za-z]{2,32}$")
AI_TRANSLATE_CACHE = {}
AI_CACHE_TTL_SECONDS = 24 * 60 * 60


def load_local_env_file():
    env_path = os.path.join(BASE_DIR, ".env")
    if not os.path.exists(env_path):
        return
    try:
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                k = k.strip()
                v = v.strip().strip('"').strip("'")
                if k and k not in os.environ:
                    os.environ[k] = v
    except Exception:
        pass


def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(OFFICIAL_DIR, exist_ok=True)
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS vocabulary (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            word TEXT NOT NULL UNIQUE,
            meaning TEXT NOT NULL,
            note TEXT DEFAULT '',
            source_book_id TEXT DEFAULT '',
            added_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS local_books (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            original_filename TEXT NOT NULL,
            stored_filename TEXT NOT NULL,
            content_text TEXT NOT NULL,
            title TEXT DEFAULT '',
            author TEXT DEFAULT '',
            source_url TEXT DEFAULT '',
            is_official INTEGER DEFAULT 0,
            added_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS books (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            gutenberg_id INTEGER UNIQUE,
            title TEXT NOT NULL,
            author TEXT NOT NULL,
            file_path TEXT NOT NULL,
            is_official INTEGER NOT NULL DEFAULT 1,
            added_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    conn.commit()
    conn.close()


def get_deepseek_client():
    load_local_env_file()
    api_key = (os.getenv("DEEPSEEK_API_KEY") or "").strip()
    if not api_key:
        return None
    return OpenAI(api_key=api_key, base_url="https://api.deepseek.com")


def deepseek_chat(messages, temperature=0.2, timeout_seconds=12):
    client = get_deepseek_client()
    if not client:
        raise RuntimeError("DEEPSEEK_API_KEY 未配置")
    resp = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages,
        temperature=temperature,
        timeout=timeout_seconds,
    )
    if not resp or not resp.choices or not resp.choices[0].message:
        return ""
    return resp.choices[0].message.content or ""


def extract_json_from_text(text):
    raw = (text or "").strip()
    if not raw:
        return {}
    try:
        return json.loads(raw)
    except Exception:
        pass
    # 兼容 markdown json code block
    import re

    m = re.search(r"```json\s*(\{.*?\})\s*```", raw, flags=re.S | re.I)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            return {}
    m2 = re.search(r"(\{.*\})", raw, flags=re.S)
    if m2:
        try:
            return json.loads(m2.group(1))
        except Exception:
            return {}
    return {}


def parse_uploaded_book(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                return f.read().strip()
    if ext == ".docx":
        from docx import Document

        doc = Document(file_path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n\n".join(parts).strip()
    if ext == ".pdf":
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        chunks = []
        for page in reader.pages:
            txt = (page.extract_text() or "").strip()
            if txt:
                chunks.append(txt)
        return "\n\n".join(chunks).strip()
    raise ValueError("unsupported file type")


def normalize_reader_text(text):
    """阅读展示前再做一层轻清洗，减少古腾堡排版噪音。"""
    import re

    content = (text or "").replace("\r\n", "\n")
    # 去掉斜体标记下划线与花括号索引
    content = content.replace("_", "")
    content = re.sub(r"\{[^\}]+\}", "", content)
    # 去掉常见插图标记/残留行
    content = re.sub(r"(?im)^.*\[(?:Illustration|Image):[^\]]*\].*\n?", "", content)
    content = re.sub(r"(?im)^[ \t]*(?:illustration|image|plate|fig\.?|figure)\b[^\n]*\n?", "", content)
    # 去掉仅由方括号组成的空残片
    content = re.sub(r"(?m)^[ \t]*[\[\]][ \t]*$", "", content)
    # 清理每行前导空白（保留最多1个空格，避免错位排版）
    lines = []
    for line in content.split("\n"):
        stripped = line.lstrip()
        if not stripped:
            lines.append("")
        else:
            lines.append(stripped)
    content = "\n".join(lines)
    # 压缩连续空行
    content = re.sub(r"\n{3,}", "\n\n", content).strip()
    return content


def mock_search(q):
    return [
        {
            "id": f"mock-{i}",
            "title": f"{q}（示例结果{i}）",
            "author": "Mock Author",
            "summary": "当前使用本地示例搜索结果。",
            "book_id": urllib.parse.quote(f"{q}-{i}", safe=""),
            "is_english": False,
        }
        for i in range(1, 4)
    ]


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/vocab")
def vocab_page():
    return render_template("vocab.html")


@app.route("/api/search")
def api_search():
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify({"source": "none", "items": []})

    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, title, author
        FROM local_books
        WHERE lower(title) LIKE lower(?)
        ORDER BY id DESC
        LIMIT 20
        """,
        (f"%{q}%",),
    ).fetchall()
    conn.close()
    if rows:
        items = [
            {
                "id": r["id"],
                "title": r["title"] or r["author"] or "Untitled",
                "author": r["author"] or "Unknown",
                "summary": "本地书库",
                "book_id": urllib.parse.quote(str(r["title"] or ""), safe=""),
                "is_english": True,
            }
            for r in rows
        ]
        return jsonify({"source": "local", "items": items})

    return jsonify({"source": "mock_fallback", "items": mock_search(q)})


@app.route("/api/recommendations")
def api_recommendations():
    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, title, author, source_url
        FROM local_books
        WHERE is_official = 1
        ORDER BY id ASC
        """
    ).fetchall()
    conn.close()
    items = [
        {
            "book_id": r["id"],
            "title": r["title"] or "Untitled",
            "cn_title": "",
            "author": r["author"] or "Unknown",
            "summary": "官方预置经典书库",
            "source_url": r["source_url"] or "",
            "is_official": True,
        }
        for r in rows
    ]
    return jsonify({"items": items})


@app.route("/api/init-official-library", methods=["POST"])
def api_init_library():
    # 从 books 表 + uploads/official 自动同步到 local_books（官方预置）
    conn = get_conn()
    cur = conn.cursor()
    inserted_or_updated = 0
    failed = []

    rows = cur.execute("SELECT gutenberg_id, title, author, file_path FROM books").fetchall()
    for r in rows:
        fp = r["file_path"] or ""
        title = r["title"] or "Untitled"
        author = r["author"] or "Unknown"
        if not fp or not os.path.exists(fp):
            failed.append({"title": title, "error": "file not found"})
            continue
        try:
            content = parse_uploaded_book(fp)
        except Exception as exc:
            failed.append({"title": title, "error": str(exc)})
            continue
        existed = cur.execute(
            "SELECT id FROM local_books WHERE is_official=1 AND lower(title)=lower(?) LIMIT 1",
            (title,),
        ).fetchone()
        original_filename = os.path.basename(fp)
        stored_filename = f"official/{original_filename}"
        source_url = f"https://www.gutenberg.org/ebooks/{r['gutenberg_id']}" if r["gutenberg_id"] else ""
        if existed:
            cur.execute(
                """
                UPDATE local_books
                SET original_filename=?, stored_filename=?, content_text=?, title=?, author=?, source_url=?, is_official=1
                WHERE id=?
                """,
                (original_filename, stored_filename, content, title, author, source_url, existed["id"]),
            )
        else:
            cur.execute(
                """
                INSERT INTO local_books (
                    original_filename, stored_filename, content_text, title, author, source_url, is_official, added_at
                ) VALUES (?, ?, ?, ?, ?, ?, 1, CURRENT_TIMESTAMP)
                """,
                (original_filename, stored_filename, content, title, author, source_url),
            )
        inserted_or_updated += 1

    # 如果 books 表为空，尝试直接扫描 uploads/official/*.txt
    if not rows and os.path.exists(OFFICIAL_DIR):
        for name in os.listdir(OFFICIAL_DIR):
            if not name.lower().endswith(".txt"):
                continue
            fp = os.path.join(OFFICIAL_DIR, name)
            title = os.path.splitext(name)[0]
            try:
                content = parse_uploaded_book(fp)
            except Exception as exc:
                failed.append({"title": title, "error": str(exc)})
                continue
            existed = cur.execute(
                "SELECT id FROM local_books WHERE is_official=1 AND lower(title)=lower(?) LIMIT 1",
                (title,),
            ).fetchone()
            stored_filename = f"official/{name}"
            if existed:
                cur.execute(
                    """
                    UPDATE local_books
                    SET original_filename=?, stored_filename=?, content_text=?, title=?, author=?, source_url=?, is_official=1
                    WHERE id=?
                    """,
                    (name, stored_filename, content, title, "Unknown", "", existed["id"]),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO local_books (
                        original_filename, stored_filename, content_text, title, author, source_url, is_official, added_at
                    ) VALUES (?, ?, ?, ?, ?, '', 1, CURRENT_TIMESTAMP)
                    """,
                    (name, stored_filename, content, title, "Unknown"),
                )
            inserted_or_updated += 1

    conn.commit()
    count = conn.execute("SELECT COUNT(*) c FROM local_books WHERE is_official = 1").fetchone()["c"]
    conn.close()
    return jsonify({"total": count, "inserted_or_updated": inserted_or_updated, "skipped": max(0, count - inserted_or_updated), "failed": failed})


@app.route("/read/<path:book_id>")
def read_book(book_id):
    decoded_title = urllib.parse.unquote(book_id)
    return render_template(
        "read.html",
        book_title=decoded_title,
        content_html="<p>该书当前使用本地阅读页。可从“经典推荐”进入官方书库阅读。</p>",
        content_text="",
        is_html=True,
        current_book_id=book_id,
    )


@app.route("/read_local/<int:book_id>")
def read_local(book_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT id, title, original_filename FROM local_books WHERE id=?",
        (book_id,),
    ).fetchone()
    conn.close()
    if not row:
        return "book not found", 404
    return render_template(
        "read.html",
        book_title=row["title"] or row["original_filename"],
        content_html="",
        content_text="",
        is_html=False,
        local_book_id=book_id,
        current_book_id=f"local:{book_id}",
    )


@app.route("/api/book/<int:book_id>")
def api_book(book_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT title, original_filename, content_text FROM local_books WHERE id=?",
        (book_id,),
    ).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "book not found"}), 404
    cleaned = normalize_reader_text(row["content_text"] or "")
    return jsonify({"title": row["title"] or row["original_filename"], "content": cleaned})


@app.route("/api/ai-translate", methods=["POST"])
def ai_translate():
    data = request.get_json(silent=True) or {}
    user_text = (data.get("word") or data.get("sentence") or "").strip()
    if not user_text:
        return jsonify({"error": "word 不能为空"}), 400
    normalized = user_text.lower()
    now = int(time.time())
    # 只对正常英文单词走缓存；短句保持实时
    if WORD_PATTERN.match(user_text):
        cached = AI_TRANSLATE_CACHE.get(normalized)
        if cached and now - cached["ts"] <= AI_CACHE_TTL_SECONDS:
            return jsonify(cached["data"])

    prompt = (
        "你是一个专业的英语助教。请解释用户输入的单词或句子，"
        "返回 JSON，字段为 meaning_cn、meaning_en、example。"
    )
    try:
        content = deepseek_chat(
            [
                {"role": "system", "content": prompt},
                {"role": "user", "content": user_text},
            ],
            temperature=0.2,
            timeout_seconds=8,
        )
        parsed = extract_json_from_text(content)
        result = {
            "meaning_cn": parsed.get("meaning_cn", "未返回中文释义"),
            "meaning_en": parsed.get("meaning_en", "未返回英文释义"),
            "example": parsed.get("example", "未返回例句"),
        }
        if WORD_PATTERN.match(user_text):
            AI_TRANSLATE_CACHE[normalized] = {"ts": now, "data": result}
        return jsonify(result)
    except Exception as exc:
        # 快速降级，避免前端长时间等待
        if WORD_PATTERN.match(user_text):
            fallback = {
                "meaning_cn": f"{user_text}（离线简释）",
                "meaning_en": "",
                "example": "",
            }
            return jsonify(fallback)
        return jsonify({"error": f"AI 翻译失败: {exc}"}), 500


@app.route("/api/ai-analyze", methods=["POST"])
def ai_analyze():
    data = request.get_json(silent=True) or {}
    sentence = (data.get("sentence") or "").strip()
    if not sentence:
        return jsonify({"error": "sentence 不能为空"}), 400
    prompt = "分析这句话的语法结构、核心短语和潜在学习点，用简洁中文回答。"
    try:
        content = deepseek_chat(
            [
                {"role": "system", "content": prompt},
                {"role": "user", "content": sentence},
            ],
            temperature=0.3,
        )
        return jsonify({"analysis": content.strip() or "未返回分析结果"})
    except Exception as exc:
        return jsonify({"error": f"AI 语法分析失败: {exc}"}), 500


@app.route("/api/vocabulary", methods=["GET"])
def api_vocabulary_list():
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, word, meaning, note, source_book_id, added_at FROM vocabulary ORDER BY id DESC"
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/vocabulary", methods=["POST"])
def api_vocabulary_create():
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    meaning = (data.get("meaning") or "").strip()
    note = (data.get("note") or "").strip()
    if not word or not meaning:
        return jsonify({"error": "word 和 meaning 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO vocabulary (word, meaning, note, added_at) VALUES (?, ?, ?, ?)",
            (word, meaning, note, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        )
        conn.commit()
        new_id = cur.lastrowid
        conn.close()
        return jsonify({"id": new_id, "word": word, "meaning": meaning, "note": note}), 201
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({"error": "该单词已存在"}), 409


@app.route("/api/vocabulary/<int:item_id>", methods=["PUT"])
def api_vocabulary_update(item_id):
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    meaning = (data.get("meaning") or "").strip()
    note = (data.get("note") or "").strip()
    if not word or not meaning:
        return jsonify({"error": "word 和 meaning 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "UPDATE vocabulary SET word=?, meaning=?, note=? WHERE id=?",
        (word, meaning, note, item_id),
    )
    conn.commit()
    changed = cur.rowcount
    conn.close()
    if changed == 0:
        return jsonify({"error": "记录不存在"}), 404
    return jsonify({"id": item_id, "word": word, "meaning": meaning, "note": note})


@app.route("/api/vocabulary/<int:item_id>", methods=["DELETE"])
def api_vocabulary_delete(item_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM vocabulary WHERE id=?", (item_id,))
    conn.commit()
    changed = cur.rowcount
    conn.close()
    if changed == 0:
        return jsonify({"error": "记录不存在"}), 404
    return jsonify({"ok": True, "deleted_id": item_id})


@app.route("/get_vocab")
def get_vocab():
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, word, meaning, added_at FROM vocabulary ORDER BY id DESC"
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/add_vocab", methods=["POST"])
def add_vocab():
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    meaning = (data.get("meaning") or "未找到释义").strip()
    if not word:
        return jsonify({"error": "word 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO vocabulary (word, meaning, note, added_at) VALUES (?, ?, '', ?)",
            (word, meaning, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        )
        conn.commit()
        new_id = cur.lastrowid
        conn.close()
        return jsonify({"ok": True, "id": new_id, "word": word, "meaning": meaning, "deduplicated": False})
    except sqlite3.IntegrityError:
        row = cur.execute(
            "SELECT id, word, meaning, added_at FROM vocabulary WHERE lower(word)=lower(?)",
            (word,),
        ).fetchone()
        conn.close()
        return jsonify(
            {
                "ok": True,
                "id": row["id"] if row else None,
                "word": row["word"] if row else word,
                "meaning": row["meaning"] if row else meaning,
                "added_at": row["added_at"] if row else None,
                "deduplicated": True,
            }
        )


@app.route("/delete_vocab/<int:item_id>", methods=["DELETE"])
def delete_vocab(item_id):
    return api_vocabulary_delete(item_id)


@app.route("/api/vocabulary/check")
def vocabulary_check():
    word = (request.args.get("word") or "").strip().lower()
    if not word:
        return jsonify({"exists": False})
    conn = get_conn()
    row = conn.execute("SELECT id FROM vocabulary WHERE lower(word)=lower(?) LIMIT 1", (word,)).fetchone()
    conn.close()
    return jsonify({"exists": bool(row)})


@app.route("/api/add_to_vocabulary", methods=["POST"])
def add_to_vocabulary():
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    definition = (data.get("definition") or "").strip() or "未找到释义"
    book_id = str(data.get("book_id") or "").strip()
    if not word:
        return jsonify({"status": "error", "error": "word 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    existed = cur.execute(
        "SELECT id FROM vocabulary WHERE lower(word)=lower(?) LIMIT 1", (word,)
    ).fetchone()
    if existed:
        conn.close()
        return jsonify({"status": "duplicate", "id": existed["id"]}), 200
    cur.execute(
        "INSERT INTO vocabulary (word, meaning, note, source_book_id, added_at) VALUES (?, ?, '', ?, ?)",
        (word, definition, book_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    )
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    return jsonify({"status": "success", "id": new_id}), 201


@app.route("/api/ecdict/highlight", methods=["POST"])
def ecdict_highlight():
    # 兼容现有前端，不阻塞页面
    return jsonify({"ecdict_ready": False, "unknown_words": []})


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "未检测到上传文件"}), 400
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"error": "文件名为空"}), 400
    original = os.path.basename(f.filename)
    ext = os.path.splitext(original)[1].lower()
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        return jsonify({"error": "仅支持 .txt/.docx/.pdf"}), 400

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    stored = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(UPLOAD_DIR, stored)
    f.save(save_path)
    try:
        content = parse_uploaded_book(save_path)
    except Exception as exc:
        return jsonify({"error": f"文件解析失败: {exc}"}), 500
    if not content:
        content = "未能提取到正文内容。"

    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO local_books (
            original_filename, stored_filename, content_text,
            title, author, source_url, is_official, added_at
        ) VALUES (?, ?, ?, ?, '', '', 0, CURRENT_TIMESTAMP)
        """,
        (original, stored, content, original),
    )
    conn.commit()
    book_id = cur.lastrowid
    conn.close()
    return jsonify({"book_id": book_id, "title": original})


@app.route("/themes")
def themes_page():
    return send_from_directory(os.path.join(BASE_DIR, "static", "themes"), "index.html")


@app.route("/themes/<path:path>")
def themes_assets(path):
    return send_from_directory(os.path.join(BASE_DIR, "static", "themes"), path)


if __name__ == "__main__":
    load_local_env_file()
    init_db()
    port = int(os.getenv("PORT", "5000"))
    debug_mode = os.getenv("FLASK_DEBUG", "0") == "1"
    app.run(host="0.0.0.0", port=port, debug=debug_mode)
=======
import json
import os
import uuid
import sqlite3
import urllib.parse
from datetime import datetime

import requests
from flask import Flask, jsonify, render_template, request
from openai import OpenAI

app = Flask(__name__)

BASE_DIR = os.path.dirname(__file__)
DB_PATH = os.path.join(BASE_DIR, "books.db")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OFFICIAL_DIR = os.path.join(UPLOAD_DIR, "official")
ALLOWED_UPLOAD_EXTENSIONS = {".txt", ".docx", ".pdf"}


def load_local_env_file():
    env_path = os.path.join(BASE_DIR, ".env")
    if not os.path.exists(env_path):
        return
    try:
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                k = k.strip()
                v = v.strip().strip('"').strip("'")
                if k and k not in os.environ:
                    os.environ[k] = v
    except Exception:
        pass


def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(OFFICIAL_DIR, exist_ok=True)
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS vocabulary (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            word TEXT NOT NULL UNIQUE,
            meaning TEXT NOT NULL,
            note TEXT DEFAULT '',
            source_book_id TEXT DEFAULT '',
            added_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS local_books (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            original_filename TEXT NOT NULL,
            stored_filename TEXT NOT NULL,
            content_text TEXT NOT NULL,
            title TEXT DEFAULT '',
            author TEXT DEFAULT '',
            source_url TEXT DEFAULT '',
            is_official INTEGER DEFAULT 0,
            added_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS books (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            gutenberg_id INTEGER UNIQUE,
            title TEXT NOT NULL,
            author TEXT NOT NULL,
            file_path TEXT NOT NULL,
            is_official INTEGER NOT NULL DEFAULT 1,
            added_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    conn.commit()
    conn.close()


def get_deepseek_client():
    load_local_env_file()
    api_key = (os.getenv("DEEPSEEK_API_KEY") or "").strip()
    if not api_key:
        return None
    return OpenAI(api_key=api_key, base_url="https://api.deepseek.com")


def deepseek_chat(messages, temperature=0.2):
    client = get_deepseek_client()
    if not client:
        raise RuntimeError("DEEPSEEK_API_KEY 未配置")
    resp = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages,
        temperature=temperature,
    )
    if not resp or not resp.choices or not resp.choices[0].message:
        return ""
    return resp.choices[0].message.content or ""


def extract_json_from_text(text):
    raw = (text or "").strip()
    if not raw:
        return {}
    try:
        return json.loads(raw)
    except Exception:
        pass
    # 兼容 markdown json code block
    import re

    m = re.search(r"```json\s*(\{.*?\})\s*```", raw, flags=re.S | re.I)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            return {}
    m2 = re.search(r"(\{.*\})", raw, flags=re.S)
    if m2:
        try:
            return json.loads(m2.group(1))
        except Exception:
            return {}
    return {}


def parse_uploaded_book(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                return f.read().strip()
    if ext == ".docx":
        from docx import Document

        doc = Document(file_path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n\n".join(parts).strip()
    if ext == ".pdf":
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        chunks = []
        for page in reader.pages:
            txt = (page.extract_text() or "").strip()
            if txt:
                chunks.append(txt)
        return "\n\n".join(chunks).strip()
    raise ValueError("unsupported file type")


def mock_search(q):
    return [
        {
            "id": f"mock-{i}",
            "title": f"{q}（示例结果{i}）",
            "author": "Mock Author",
            "summary": "当前使用本地示例搜索结果。",
            "book_id": urllib.parse.quote(f"{q}-{i}", safe=""),
            "is_english": False,
        }
        for i in range(1, 4)
    ]


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/vocab")
def vocab_page():
    return render_template("vocab.html")


@app.route("/api/search")
def api_search():
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify({"source": "none", "items": []})

    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, title, author
        FROM local_books
        WHERE lower(title) LIKE lower(?)
        ORDER BY id DESC
        LIMIT 20
        """,
        (f"%{q}%",),
    ).fetchall()
    conn.close()
    if rows:
        items = [
            {
                "id": r["id"],
                "title": r["title"] or r["author"] or "Untitled",
                "author": r["author"] or "Unknown",
                "summary": "本地书库",
                "book_id": urllib.parse.quote(str(r["title"] or ""), safe=""),
                "is_english": True,
            }
            for r in rows
        ]
        return jsonify({"source": "local", "items": items})

    return jsonify({"source": "mock_fallback", "items": mock_search(q)})


@app.route("/api/recommendations")
def api_recommendations():
    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, title, author, source_url
        FROM local_books
        WHERE is_official = 1
        ORDER BY id ASC
        """
    ).fetchall()
    conn.close()
    items = [
        {
            "book_id": r["id"],
            "title": r["title"] or "Untitled",
            "cn_title": "",
            "author": r["author"] or "Unknown",
            "summary": "官方预置经典书库",
            "source_url": r["source_url"] or "",
            "is_official": True,
        }
        for r in rows
    ]
    return jsonify({"items": items})


@app.route("/api/init-official-library", methods=["POST"])
def api_init_library():
    # 从 books 表 + uploads/official 自动同步到 local_books（官方预置）
    conn = get_conn()
    cur = conn.cursor()
    inserted_or_updated = 0
    failed = []

    rows = cur.execute("SELECT gutenberg_id, title, author, file_path FROM books").fetchall()
    for r in rows:
        fp = r["file_path"] or ""
        title = r["title"] or "Untitled"
        author = r["author"] or "Unknown"
        if not fp or not os.path.exists(fp):
            failed.append({"title": title, "error": "file not found"})
            continue
        try:
            content = parse_uploaded_book(fp)
        except Exception as exc:
            failed.append({"title": title, "error": str(exc)})
            continue
        existed = cur.execute(
            "SELECT id FROM local_books WHERE is_official=1 AND lower(title)=lower(?) LIMIT 1",
            (title,),
        ).fetchone()
        original_filename = os.path.basename(fp)
        stored_filename = f"official/{original_filename}"
        source_url = f"https://www.gutenberg.org/ebooks/{r['gutenberg_id']}" if r["gutenberg_id"] else ""
        if existed:
            cur.execute(
                """
                UPDATE local_books
                SET original_filename=?, stored_filename=?, content_text=?, title=?, author=?, source_url=?, is_official=1
                WHERE id=?
                """,
                (original_filename, stored_filename, content, title, author, source_url, existed["id"]),
            )
        else:
            cur.execute(
                """
                INSERT INTO local_books (
                    original_filename, stored_filename, content_text, title, author, source_url, is_official, added_at
                ) VALUES (?, ?, ?, ?, ?, ?, 1, CURRENT_TIMESTAMP)
                """,
                (original_filename, stored_filename, content, title, author, source_url),
            )
        inserted_or_updated += 1

    # 如果 books 表为空，尝试直接扫描 uploads/official/*.txt
    if not rows and os.path.exists(OFFICIAL_DIR):
        for name in os.listdir(OFFICIAL_DIR):
            if not name.lower().endswith(".txt"):
                continue
            fp = os.path.join(OFFICIAL_DIR, name)
            title = os.path.splitext(name)[0]
            try:
                content = parse_uploaded_book(fp)
            except Exception as exc:
                failed.append({"title": title, "error": str(exc)})
                continue
            existed = cur.execute(
                "SELECT id FROM local_books WHERE is_official=1 AND lower(title)=lower(?) LIMIT 1",
                (title,),
            ).fetchone()
            stored_filename = f"official/{name}"
            if existed:
                cur.execute(
                    """
                    UPDATE local_books
                    SET original_filename=?, stored_filename=?, content_text=?, title=?, author=?, source_url=?, is_official=1
                    WHERE id=?
                    """,
                    (name, stored_filename, content, title, "Unknown", "", existed["id"]),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO local_books (
                        original_filename, stored_filename, content_text, title, author, source_url, is_official, added_at
                    ) VALUES (?, ?, ?, ?, ?, '', 1, CURRENT_TIMESTAMP)
                    """,
                    (name, stored_filename, content, title, "Unknown"),
                )
            inserted_or_updated += 1

    conn.commit()
    count = conn.execute("SELECT COUNT(*) c FROM local_books WHERE is_official = 1").fetchone()["c"]
    conn.close()
    return jsonify({"total": count, "inserted_or_updated": inserted_or_updated, "skipped": max(0, count - inserted_or_updated), "failed": failed})


@app.route("/read/<path:book_id>")
def read_book(book_id):
    decoded_title = urllib.parse.unquote(book_id)
    return render_template(
        "read.html",
        book_title=decoded_title,
        content_html="<p>该书当前使用本地阅读页。可从“经典推荐”进入官方书库阅读。</p>",
        content_text="",
        is_html=True,
        current_book_id=book_id,
    )


@app.route("/read_local/<int:book_id>")
def read_local(book_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT id, title, original_filename FROM local_books WHERE id=?",
        (book_id,),
    ).fetchone()
    conn.close()
    if not row:
        return "book not found", 404
    return render_template(
        "read.html",
        book_title=row["title"] or row["original_filename"],
        content_html="",
        content_text="",
        is_html=False,
        local_book_id=book_id,
        current_book_id=f"local:{book_id}",
    )


@app.route("/api/book/<int:book_id>")
def api_book(book_id):
    conn = get_conn()
    row = conn.execute(
        "SELECT title, original_filename, content_text FROM local_books WHERE id=?",
        (book_id,),
    ).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "book not found"}), 404
    return jsonify({"title": row["title"] or row["original_filename"], "content": row["content_text"] or ""})


@app.route("/api/ai-translate", methods=["POST"])
def ai_translate():
    data = request.get_json(silent=True) or {}
    user_text = (data.get("word") or data.get("sentence") or "").strip()
    if not user_text:
        return jsonify({"error": "word 不能为空"}), 400
    prompt = (
        "你是一个专业的英语助教。请解释用户输入的单词或句子，"
        "返回 JSON，字段为 meaning_cn、meaning_en、example。"
    )
    try:
        content = deepseek_chat(
            [
                {"role": "system", "content": prompt},
                {"role": "user", "content": user_text},
            ],
            temperature=0.2,
        )
        parsed = extract_json_from_text(content)
        return jsonify(
            {
                "meaning_cn": parsed.get("meaning_cn", "未返回中文释义"),
                "meaning_en": parsed.get("meaning_en", "未返回英文释义"),
                "example": parsed.get("example", "未返回例句"),
            }
        )
    except Exception as exc:
        return jsonify({"error": f"AI 翻译失败: {exc}"}), 500


@app.route("/api/ai-analyze", methods=["POST"])
def ai_analyze():
    data = request.get_json(silent=True) or {}
    sentence = (data.get("sentence") or "").strip()
    if not sentence:
        return jsonify({"error": "sentence 不能为空"}), 400
    prompt = "分析这句话的语法结构、核心短语和潜在学习点，用简洁中文回答。"
    try:
        content = deepseek_chat(
            [
                {"role": "system", "content": prompt},
                {"role": "user", "content": sentence},
            ],
            temperature=0.3,
        )
        return jsonify({"analysis": content.strip() or "未返回分析结果"})
    except Exception as exc:
        return jsonify({"error": f"AI 语法分析失败: {exc}"}), 500


@app.route("/api/vocabulary", methods=["GET"])
def api_vocabulary_list():
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, word, meaning, note, source_book_id, added_at FROM vocabulary ORDER BY id DESC"
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/vocabulary", methods=["POST"])
def api_vocabulary_create():
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    meaning = (data.get("meaning") or "").strip()
    note = (data.get("note") or "").strip()
    if not word or not meaning:
        return jsonify({"error": "word 和 meaning 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO vocabulary (word, meaning, note, added_at) VALUES (?, ?, ?, ?)",
            (word, meaning, note, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        )
        conn.commit()
        new_id = cur.lastrowid
        conn.close()
        return jsonify({"id": new_id, "word": word, "meaning": meaning, "note": note}), 201
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({"error": "该单词已存在"}), 409


@app.route("/api/vocabulary/<int:item_id>", methods=["PUT"])
def api_vocabulary_update(item_id):
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    meaning = (data.get("meaning") or "").strip()
    note = (data.get("note") or "").strip()
    if not word or not meaning:
        return jsonify({"error": "word 和 meaning 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "UPDATE vocabulary SET word=?, meaning=?, note=? WHERE id=?",
        (word, meaning, note, item_id),
    )
    conn.commit()
    changed = cur.rowcount
    conn.close()
    if changed == 0:
        return jsonify({"error": "记录不存在"}), 404
    return jsonify({"id": item_id, "word": word, "meaning": meaning, "note": note})


@app.route("/api/vocabulary/<int:item_id>", methods=["DELETE"])
def api_vocabulary_delete(item_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM vocabulary WHERE id=?", (item_id,))
    conn.commit()
    changed = cur.rowcount
    conn.close()
    if changed == 0:
        return jsonify({"error": "记录不存在"}), 404
    return jsonify({"ok": True, "deleted_id": item_id})


@app.route("/get_vocab")
def get_vocab():
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, word, meaning, added_at FROM vocabulary ORDER BY id DESC"
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/add_vocab", methods=["POST"])
def add_vocab():
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    meaning = (data.get("meaning") or "未找到释义").strip()
    if not word:
        return jsonify({"error": "word 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO vocabulary (word, meaning, note, added_at) VALUES (?, ?, '', ?)",
            (word, meaning, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        )
        conn.commit()
        new_id = cur.lastrowid
        conn.close()
        return jsonify({"ok": True, "id": new_id, "word": word, "meaning": meaning, "deduplicated": False})
    except sqlite3.IntegrityError:
        row = cur.execute(
            "SELECT id, word, meaning, added_at FROM vocabulary WHERE lower(word)=lower(?)",
            (word,),
        ).fetchone()
        conn.close()
        return jsonify(
            {
                "ok": True,
                "id": row["id"] if row else None,
                "word": row["word"] if row else word,
                "meaning": row["meaning"] if row else meaning,
                "added_at": row["added_at"] if row else None,
                "deduplicated": True,
            }
        )


@app.route("/delete_vocab/<int:item_id>", methods=["DELETE"])
def delete_vocab(item_id):
    return api_vocabulary_delete(item_id)


@app.route("/api/vocabulary/check")
def vocabulary_check():
    word = (request.args.get("word") or "").strip().lower()
    if not word:
        return jsonify({"exists": False})
    conn = get_conn()
    row = conn.execute("SELECT id FROM vocabulary WHERE lower(word)=lower(?) LIMIT 1", (word,)).fetchone()
    conn.close()
    return jsonify({"exists": bool(row)})


@app.route("/api/add_to_vocabulary", methods=["POST"])
def add_to_vocabulary():
    data = request.get_json(silent=True) or {}
    word = (data.get("word") or "").strip()
    definition = (data.get("definition") or "").strip() or "未找到释义"
    book_id = str(data.get("book_id") or "").strip()
    if not word:
        return jsonify({"status": "error", "error": "word 不能为空"}), 400
    conn = get_conn()
    cur = conn.cursor()
    existed = cur.execute(
        "SELECT id FROM vocabulary WHERE lower(word)=lower(?) LIMIT 1", (word,)
    ).fetchone()
    if existed:
        conn.close()
        return jsonify({"status": "duplicate", "id": existed["id"]}), 200
    cur.execute(
        "INSERT INTO vocabulary (word, meaning, note, source_book_id, added_at) VALUES (?, ?, '', ?, ?)",
        (word, definition, book_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    )
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    return jsonify({"status": "success", "id": new_id}), 201


@app.route("/api/ecdict/highlight", methods=["POST"])
def ecdict_highlight():
    # 兼容现有前端，不阻塞页面
    return jsonify({"ecdict_ready": False, "unknown_words": []})


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "未检测到上传文件"}), 400
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"error": "文件名为空"}), 400
    original = os.path.basename(f.filename)
    ext = os.path.splitext(original)[1].lower()
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        return jsonify({"error": "仅支持 .txt/.docx/.pdf"}), 400

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    stored = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(UPLOAD_DIR, stored)
    f.save(save_path)
    try:
        content = parse_uploaded_book(save_path)
    except Exception as exc:
        return jsonify({"error": f"文件解析失败: {exc}"}), 500
    if not content:
        content = "未能提取到正文内容。"

    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO local_books (
            original_filename, stored_filename, content_text,
            title, author, source_url, is_official, added_at
        ) VALUES (?, ?, ?, ?, '', '', 0, CURRENT_TIMESTAMP)
        """,
        (original, stored, content, original),
    )
    conn.commit()
    book_id = cur.lastrowid
    conn.close()
    return jsonify({"book_id": book_id, "title": original})


if __name__ == "__main__":
    load_local_env_file()
    init_db()
    port = int(os.getenv("PORT", "5000"))
    debug_mode = os.getenv("FLASK_DEBUG", "0") == "1"
    app.run(host="0.0.0.0", port=port, debug=debug_mode)
>>>>>>> d490cee180814eecf12c3a91283f9c9719b4abe6
